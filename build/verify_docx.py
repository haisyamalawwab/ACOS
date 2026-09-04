"""Verify the DOCX carries the same content as the parsed source, with no markdown left."""
from __future__ import annotations

import json
import os
import re
import sys

import docx

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCX = os.path.join(ROOT, "dist", "Panduan_Anotasi_ACOSE_Bahasa_Indonesia.docx")
JSON = os.path.join(os.path.dirname(os.path.abspath(__file__)), "book.json")

book = json.load(open(JSON, encoding="utf-8"))
doc = docx.Document(DOCX)


def cell_texts(table) -> list[str]:
    return [c.text for row in table.rows for c in row.cells]


all_para = [p.text for p in doc.paragraphs]
all_cell = [t for tb in doc.tables for t in cell_texts(tb)]
haystack = re.sub(r"\s+", " ", " ".join(all_para + all_cell))

print("paragraphs:", len(all_para), "| tables:", len(doc.tables))

styles = {}
for p in doc.paragraphs:
    name = p.style.name if p.style is not None else "(none)"
    styles[name] = styles.get(name, 0) + 1
print("heading paragraphs:",
      {k: v for k, v in sorted(styles.items()) if k.startswith("Heading")})

print()
print("=== chapter titles present ===")
missing = [ch["title"] for ch in book["chapters"]
           if re.sub(r"\s+", " ", ch["title"]) not in haystack]
print("   ", "ALL 9 PRESENT" if not missing else missing)

print("=== source headings present ===")
gone = []
for ch in book["chapters"]:
    for b in ch["blocks"]:
        if b["kind"] != "heading" or b.get("flag") == "tech":
            continue
        txt = re.sub(r"[*`]", "", b["text"])
        if re.sub(r"\s+", " ", txt)[:40] not in haystack:
            gone.append((ch["stem"], b["text"][:45]))
print("   ", "ALL PRESENT" if not gone else gone[:6])

print("=== paragraph text present (sampled) ===")
lost = []
for ch in book["chapters"]:
    paras = [b for b in ch["blocks"] if b["kind"] == "para"]
    for b in paras[:: max(len(paras) // 4, 1)]:
        probe = re.sub(r"[*`]", "", b["text"])
        probe = re.sub(r"\s+", " ", probe)[:45]
        if probe and probe not in haystack:
            lost.append((ch["stem"], probe))
print("   ", "ALL PRESENT" if not lost else lost[:6])

print("=== markdown residue ===")
for label, pattern in (
    ("bold **", r"\*\*"),
    ("heading ##", r"(?m)^#{1,4}\s"),
    ("backticks", r"`"),
    ("pipe table", r"\|\s*-{3,}"),
    ("md links", r"\]\(\./"),
    ("file codes 022x", r"\b022(aa|[a-h])\b"),
):
    hits = re.findall(pattern, haystack)
    print("    %-18s %s" % (label, "CLEAN" if not hits else "FOUND %d" % len(hits)))

print("=== structural blocks ===")
want = {"table": 0, "code": 0, "quote": 0}
for ch in book["chapters"]:
    for b in ch["blocks"]:
        if b["kind"] in want:
            want[b["kind"]] += 1
tech = sum(1 for ch in book["chapters"] for b in ch["blocks"]
           if b["kind"] == "heading" and b.get("flag") == "tech")
expected_tables = want["table"] + want["code"] + want["quote"] + tech
print("    source: %d data tables + %d code + %d callouts + %d tech bands = %d table objects"
      % (want["table"], want["code"], want["quote"], tech, expected_tables))
print("    docx  : %d table objects (1 extra = cover wrapper) -> %s"
      % (len(doc.tables), "OK" if len(doc.tables) == expected_tables + 1 else "MISMATCH"))

print("=== sections & fields ===")
print("    sections:", len(doc.sections))
import zipfile

with zipfile.ZipFile(DOCX) as z:
    names = z.namelist()
    headers = sorted(n for n in names if re.match(r"word/header\d+\.xml", n))
    footers = sorted(n for n in names if re.match(r"word/footer\d+\.xml", n))
    header_texts = []
    for n in headers:
        xml = z.read(n).decode("utf-8")
        header_texts.append(" | ".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)))
    footer_xml = "".join(z.read(n).decode("utf-8") for n in footers)
print("    header parts:", len(headers), "| footer parts:", len(footers))
print("    footer format switches:",
      sorted(set(re.findall(r"PAGE \\\* (\w+)", footer_xml))))
body_xml = doc.element.xml
print("    TOC field:", "yes" if re.search(r"TOC\b", body_xml) else "NO")
print("    TOC placeholder entries:", body_xml.count("PAGEREF"))
print("    empty pgNumType left:", body_xml.count("<w:pgNumType/>"))

print("=== running head per chapter ===")
want = [ch["running"] for ch in book["chapters"]]
got = [t for t in header_texts if t.strip()]
for i, text in enumerate(got):
    tail = text.split("|")[-1].strip().replace("&amp;", "&")
    mark = "OK" if tail in want else "?"
    print("    header%-2d %-22s %s" % (i + 1, tail, mark))
print("    chapters covered: %d/%d" % (len(got), len(want)))
